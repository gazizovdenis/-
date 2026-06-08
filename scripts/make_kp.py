import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2)
section.right_margin = Cm(1.5)
section.page_width = Cm(21)
section.page_height = Cm(29.7)

def shd_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        tblBorders.append(el)
    tblPr.append(tblBorders)

def fmt(n):
    parts = f'{n:.2f}'.split('.')
    s = parts[0]
    r = ''
    for i, c in enumerate(reversed(s)):
        if i > 0 and i % 3 == 0:
            r = ' ' + r
        r = c + r
    return r + ',' + parts[1]

# === ДАННЫЕ ===
items = [
    {
        'num': 1,
        'name': 'Накладки держателей бутылок UNIBLOC MOD VACUUM 40/40/8',
        'qty': 82,
        'price': round(3850 * 1.1, 2),
    },
    {
        'num': 2,
        'name': 'Накладки держателей бутылок инспекционной машины MOD.VISION 24',
        'qty': 98,
        'price': round(3850 * 1.1, 2),
    },
    {
        'num': 3,
        'name': 'Столы донной ориентации для этикетировочного автомата Z-ADHESIVE 96/8T',
        'qty': 17,
        'price': round(12000 * 1.1, 2),
    },
    {
        'num': 4,
        'name': 'Полный комплект оснастки системы подачи колпачков укупорочного автомата',
        'qty': 1,
        'price': round(418000 * 1.1, 2),
    },
    {
        'num': 5,
        'name': 'Оснастка на укупорочный автомат (звездочка входная, рассекатель, звездочка выходная, шнек, кольцевая направляющая)',
        'qty': 1,
        'price': round(385000 * 1.1, 2),
    },
]

for item in items:
    item['sum'] = round(item['qty'] * item['price'], 2)

total_no_vat = round(sum(i['sum'] for i in items), 2)
vat = round(total_no_vat * 0.22, 2)
total_with_vat = round(total_no_vat + vat, 2)

# === ШАПКА (таблица 2 строки x 2 столбца) ===
htbl = doc.add_table(rows=2, cols=2)
htbl.style = 'Table Grid'

# Строка 0: логотип | реквизиты
cl = htbl.cell(0, 0)
shd_cell(cl, 'FFFFFF')
p = cl.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run('TECHNSPRODUCT')
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0, 100, 0)
p2 = cl.add_paragraph(
    'Импорт промышленной автоматики, электроники и з/ч к\n'
    'оборудованию. Изготовление по чертежам. Услуги\n'
    'аккредитованного агента по закупкам'
)
p2.runs[0].italic = True
p2.runs[0].font.size = Pt(8)

cr = htbl.cell(0, 1)
shd_cell(cr, 'FFFFFF')
p = cr.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('ООО «Технопродукт»')
r.bold = True
r.font.size = Pt(11)
cr.add_paragraph()
p2 = cr.add_paragraph('ИНН/КПП: 0275914810/027501001')
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p2.runs[0].bold = True
p2.runs[0].font.size = Pt(10)

# Строка 1: номер/дата | получатель
cl2 = htbl.cell(1, 0)
shd_cell(cl2, 'F2F2F2')
p = cl2.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Исх. №250 от 25.05.2026')
r.font.size = Pt(10)

cr2 = htbl.cell(1, 1)
shd_cell(cr2, 'F2F2F2')
p = cr2.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Для: АО "БАШСПИРТ"')
r.bold = True
r.italic = True
r.font.size = Pt(11)

set_borders(htbl)
doc.add_paragraph()

# === ЗАГОЛОВОК ===
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Коммерческое предложение')
r.bold = True
r.font.size = Pt(16)
doc.add_paragraph()

# === ВВОДНЫЙ ТЕКСТ ===
intro = doc.add_paragraph(
    'Выражаем свою признательность за интерес к нашей компании, '
    'и сообщаем о возможности поставки продукции по следующим ценам:'
)
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
intro.runs[0].font.size = Pt(11)
doc.add_paragraph()

# === ОСНОВНАЯ ТАБЛИЦА ===
# 6 столбцов: №, Наименование, Кол-во, Цена без НДС, Сумма без НДС, Срок
num_rows = 1 + len(items) + 3  # заголовок + строки + итоги
mtbl = doc.add_table(rows=num_rows, cols=6)
mtbl.style = 'Table Grid'

col_widths = [Cm(0.9), Cm(5.8), Cm(1.4), Cm(2.5), Cm(2.5), Cm(1.5)]
for ri in range(num_rows):
    for ci, w in enumerate(col_widths):
        mtbl.rows[ri].cells[ci].width = w

# Заголовок таблицы
headers = [
    'No.',
    'Наименование',
    'Кол-во',
    'Цена, руб\nбез НДС',
    'Сумма, руб\nбез НДС',
    'Срок\nпоставки\n(дней)',
]
for ci, h in enumerate(headers):
    c = mtbl.cell(0, ci)
    shd_cell(c, 'D9D9D9')
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(9)

# Строки данных
col_aligns = [
    WD_ALIGN_PARAGRAPH.CENTER,
    WD_ALIGN_PARAGRAPH.LEFT,
    WD_ALIGN_PARAGRAPH.CENTER,
    WD_ALIGN_PARAGRAPH.RIGHT,
    WD_ALIGN_PARAGRAPH.RIGHT,
    WD_ALIGN_PARAGRAPH.CENTER,
]
for ri, item in enumerate(items):
    row_idx = ri + 1
    vals = [
        str(item['num']),
        item['name'],
        str(item['qty']),
        fmt(item['price']),
        fmt(item['sum']),
        '',
    ]
    for ci, (v, al) in enumerate(zip(vals, col_aligns)):
        c = mtbl.cell(row_idx, ci)
        p = c.paragraphs[0]
        p.alignment = al
        run = p.add_run(v)
        run.font.size = Pt(9)

# Итоговые строки
summary_start = len(items) + 1
totals = [
    ('Сумма без НДС', fmt(total_no_vat)),
    ('НДС 22%', fmt(vat)),
    ('Сумма с НДС 22%', fmt(total_with_vat)),
]
for i, (label, val) in enumerate(totals):
    r_idx = summary_start + i
    merged = mtbl.cell(r_idx, 0).merge(mtbl.cell(r_idx, 4))
    p = merged.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(9)
    vc = mtbl.cell(r_idx, 5)
    p2 = vc.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(val)
    r2.bold = True
    r2.font.size = Pt(9)

set_borders(mtbl)
doc.add_paragraph()

# === УСЛОВИЯ ===
for txt in [
    'Цены указаны с учетом стоимости доставки.',
    'Условия оплаты: 50% предоплата, 50% по факту поставки',
    'Срок поставки: До 14 дней',
]:
    p = doc.add_paragraph(txt)
    p.runs[0].font.size = Pt(10)

doc.add_paragraph()

# === ИСПОЛНИТЕЛЬ ===
p = doc.add_paragraph('Исполнитель: Салават Шигабутдинов  +7 960 384 00 16')
p.runs[0].font.size = Pt(10)
p = doc.add_paragraph('e-mail: sh.salavat@tehnoprod.ru')
p.runs[0].font.size = Pt(10)
doc.add_paragraph()

# === ДИРЕКТОР ===
p = doc.add_paragraph()
p.add_run('Директор').font.size = Pt(10)
p.add_run('\t\t\t').font.size = Pt(10)
p.add_run('Газизов Д.Р').font.size = Pt(10)
doc.add_paragraph()

# === ПОДВАЛ ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Оперативно обработаем ваши заявки')
r.bold = True
r.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('эл.почта: sh.salavat@tehnoprod.ru')
r.bold = True
r.font.size = Pt(11)
doc.add_paragraph()

# === ПРЕИМУЩЕСТВА ===
p = doc.add_paragraph()
r = p.add_run('Технопродукт ')
r.bold = True
r.font.size = Pt(9)

advantages = [
    'НАДЕЖНОСТЬ: более 5 лет на рынке, долгосрочные контракты с крупнейшими компаниями РФ',
    'ПЕРСОНАЛЬНОЕ ОБСЛУЖИВАНИЕ: закреплённый менеджер ведёт заявку от приёма до отгрузки',
    'СРОК ОТВЕТА 2-3 ДНЯ: налаженный опыт с зарубежными производителями и поставщиками',
    'ОПТИМАЛЬНЫЕ ЦЕНЫ: оптовые закупки позволяют сокращать издержки логистики и платежей',
]
for adv in advantages:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(adv)
    run.font.size = Pt(8)

# === СОХРАНЕНИЕ ===
out_path = r'Хранилище/raw/02_Технопродукт/КП/КП_на_основе_КП2_цена_плюс_10.docx'
doc.save(out_path)

print(f'Файл сохранён: {out_path}')
print(f'Позиций перенесено: {len(items)}')
print(f'Цена увеличена на 10%: да')
print(f'Сумма без НДС: {fmt(total_no_vat)} руб.')
print(f'НДС 22%: {fmt(vat)} руб.')
print(f'Сумма с НДС 22%: {fmt(total_with_vat)} руб.')
print()
print('Расчёт по позициям:')
for item in items:
    print(f'  {item["num"]}. {item["name"][:50]}...')
    print(f'     Кол-во: {item["qty"]}, Цена: {fmt(item["price"])}, Сумма: {fmt(item["sum"])}')
